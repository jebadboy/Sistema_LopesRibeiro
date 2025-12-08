"""
Módulo de IA Jurídica - Interface do Usuário
Sistema Lopes &  Ribeiro
"""

import streamlit as st
import ai_gemini as ai
from datetime import datetime
import database as db
import pandas as pd
import PyPDF2
import docx
from docx import Document
import io
from io import BytesIO

def extract_text_from_pdf(file):
    """Extrai texto de um arquivo PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Erro ao ler PDF: {e}"

def extract_text_from_docx(file):
    """Extrai texto de um arquivo DOCX"""
    try:
        doc = docx.Document(file)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    except Exception as e:
        return f"Erro ao ler DOCX: {e}"

def extract_text_from_txt(file):
    """Extrai texto de um arquivo TXT"""
    try:
        return file.getvalue().decode("utf-8")
    except Exception as e:
        return f"Erro ao ler TXT: {e}"

def gerar_docx(texto_ia):
    """Gera um arquivo DOCX com a resposta da IA"""
    doc = Document()
    doc.add_heading("Lopes & Ribeiro Advocacia - Parecer Jurídico", 0)
    doc.add_paragraph(texto_ia)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def render():
    """Renderiza a interface do módulo de IA Jurídica"""
    
    # Inicializar IA se ainda não foi feito
    if 'ai_inicializada' not in st.session_state:
        with st.spinner("Inicializando IA..."):
            st.session_state.ai_inicializada = ai.inicializar_gemini()
    
    if not st.session_state.ai_inicializada:
        st.warning("⚠️ IA não inicializada. Algumas funcionalidades podem estar indisponíveis. Verifique a configuração da API Gemini.")
        # Não retorna mais, permite carregar o histórico

    
    st.title("🤖 Assistente Jurídico Inteligente")
    st.caption("Powered by Google Gemini AI")
    
    # Tabs principais
    tab1, tab2, tab3, tab4 = st.tabs([
        "💬 Chat Assistente",
        "📄 Análise de Documentos",
        "💡 Sugestões Inteligentes",
        "📚 Histórico"
    ])
    
    # TAB 1: Chat Assistente
    with tab1:
        render_chat()
    
    # TAB 2: Análise de Documentos
    with tab2:
        render_analise_documentos()
    
    # TAB 3: Sugestões Inteligentes
    with tab3:
        render_sugestoes()
    
    # TAB 4: Histórico
    with tab4:
        render_historico()


def render_chat():
    """Renderiza interface de chat"""
    st.subheader("💬 Converse com o Assistente Jurídico")
    
    if not st.session_state.get('ai_inicializada'):
        st.error("❌ IA Offline. Verifique configurações.")
        return
    
    # Ações Rápidas
    render_quick_actions()
    
    # Inicializar histórico de chat
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = [
            {"role": "assistant", "content": "Entendido. Estou pronto para atuar como Consultor Sênior do Lopes & Ribeiro. Qual o próximo caso?"}
        ]
    
    # Exibir histórico de mensagens
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
    
    # Input de mensagem
    if prompt := st.chat_input("Digite sua pergunta jurídica..."):
        # Adicionar mensagem do usuário
        st.session_state.chat_history.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Gerar resposta da IA
        with st.chat_message("assistant"):
            with st.spinner("Pensando..."):
                resposta = ai.chat_assistente(prompt)
                st.markdown(resposta)
                
                # Botão de Download
                st.download_button(
                    label="📥 Baixar Parecer em Word (.docx)",
                    data=gerar_docx(resposta),
                    file_name="parecer_lopes_ribeiro_IA.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        # Adicionar resposta ao histórico
        st.session_state.chat_history.append({"role": "assistant", "content": resposta})
        
        # Salvar no banco de dados
        try:
            with db.get_connection() as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO ai_historico (usuario, tipo, input, output, data_hora)
                    VALUES (?, ?, ?, ?, ?)
                """, (st.session_state.user, 'chat', prompt, resposta, datetime.now().isoformat()))
                conn.commit()
        except Exception as e:
            st.error(f"Erro ao salvar histórico: {e}")
    
    # Botão para limpar chat
    if st.button("🗑️ Limpar Conversa"):
        st.session_state.chat_history = []
        st.rerun()


def render_analise_documentos():
    """Renderiza interface de análise de documentos"""
    st.subheader("📄 Análise Inteligente de Documentos")
    
    if not st.session_state.get('ai_inicializada'):
        st.error("❌ IA Offline.")
        return
    
    st.info("🔍 Cole o texto ou faça upload de um documento jurídico para análise automática")
    
    col1, col2 = st.columns([3, 1])
    
    with col1:
        tipo_doc = st.selectbox(
            "Tipo de Documento",
            ["Petição Inicial", "Contrato", "Sentença", "Acórdão", "Outro"]
        )
    
    # Upload de arquivo
    uploaded_file = st.file_uploader("Carregar arquivo (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'])
    
    texto_extraido = ""
    if uploaded_file is not None:
        file_type = uploaded_file.name.split('.')[-1].lower()
        if file_type == 'pdf':
            texto_extraido = extract_text_from_pdf(uploaded_file)
        elif file_type == 'docx':
            texto_extraido = extract_text_from_docx(uploaded_file)
        elif file_type == 'txt':
            texto_extraido = extract_text_from_txt(uploaded_file)
            
        if texto_extraido.startswith("Erro"):
            st.error(texto_extraido)
            texto_extraido = ""
        elif not texto_extraido.strip():
            st.warning("⚠️ O texto extraído está vazio ou ilegível. Se este documento for um PDF escaneado (imagem), a IA não conseguirá ler o conteúdo.")
        else:
            st.success(f"Arquivo '{uploaded_file.name}' carregado com sucesso!")
    
    texto_documento = st.text_area(
        "Conteúdo do documento:",
        value=texto_extraido if texto_extraido else "",
        height=300,
        placeholder="Cole aqui o conteúdo do documento ou faça upload de um arquivo..."
    )
    
    if st.button("🔍 Analisar Documento", type="primary"):
        if not texto_documento:
            st.warning("Por favor, forneça um texto para análise (cole ou faça upload)")
            return
        
        with st.spinner("Analisando documento..."):
            resultado = ai.analisar_documento(texto_documento, tipo_doc.lower())
            
            if 'erro' in resultado:
                st.error(f"❌ Erro na análise: {resultado['erro']}")
            else:
                st.success("✅ Análise concluída!")
                
                if resultado.get('from_cache'):
                    st.info("📦 Resultado obtido do cache (análise anterior)")
                
                # Exibir análise
                st.markdown("### 📊 Resultado da Análise")
                st.markdown(resultado['analise'])
                
                # Botão de Download
                st.download_button(
                    label="📥 Baixar Parecer em Word (.docx)",
                    data=gerar_docx(resultado['analise']),
                    file_name="analise_documento_lopes_ribeiro.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Salvar no histórico
                try:
                    with db.get_connection() as conn:
                        cursor = conn.cursor()
                        cursor.execute("""
                            INSERT INTO ai_historico (usuario, tipo, input, output, data_hora)
                            VALUES (?, ?, ?, ?, ?)
                        """, (st.session_state.user, 'analise', texto_documento[:500], resultado['analise'], datetime.now().isoformat()))
                        conn.commit()
                except Exception as e:
                    st.error(f"Erro ao salvar análise: {e}")


def render_sugestoes():
    """Renderiza sugestões inteligentes baseadas em processos"""
    st.subheader("💡 Sugestões Inteligentes")
    
    if not st.session_state.get('ai_inicializada'):
        st.error("❌ IA Offline.")
        return
    
    # Buscar processos ativos
    try:
        processos = db.sql_get('processos')
        
        if processos.empty:
            st.info("Nenhum processo encontrado para análise.")
            return

        # Verificar se coluna status existe (compatibilidade)
        if 'status' not in processos.columns:
            # Tentar usar status_processo ou considerar todos ativos se não tiver filtro
            if 'status_processo' in processos.columns:
                 processos_ativos = processos # Fallback, assume todos
            else:
                 processos_ativos = processos
        else:
            processos_ativos = processos[processos['status'] == 'Ativo']
        
        if processos_ativos.empty:
            st.info("Nenhum processo ativo encontrado")
            return
        
        selected_processo_id = st.selectbox(
            "Selecione um processo para obter sugestões:",
            options=processos_ativos['id'].tolist(),
            format_func=lambda x: f"Proc. {x} - {processos_ativos[processos_ativos['id']==x]['cliente_nome'].values[0]}"
        )
        
        if st.button("💡 Gerar Sugestões", type="primary"):
            # Buscar dados do processo
            processo_dados = processos_ativos[processos_ativos['id'] == selected_processo_id].iloc[0].to_dict()
            
            with st.spinner("Gerando sugestões..."):
                prompt = f"""
                Com base nos dados do processo abaixo, forneça 5 sugestões práticas de próximas ações:
                
                Número (ID): {processo_dados.get('id')}
                Ação: {processo_dados.get('acao')}
                Status: {processo_dados.get('status')}
                Observações: {processo_dados.get('obs', 'Nenhuma')}
                
                Liste as sugestões de forma objetiva e prática.
                """
                
                resposta = ai.chat_assistente(prompt, contexto=processo_dados)
                
                st.markdown("### 📋 Sugestões Geradas")
                st.markdown(resposta)
                
                # Botão de Download
                st.download_button(
                    label="📥 Baixar Sugestões em Word (.docx)",
                    data=gerar_docx(resposta),
                    file_name="sugestoes_lopes_ribeiro.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Salvar no histórico
                try:
                    with db.get_connection() as conn:
                        cursor = conn.cursor()
                        processo_id = processo_dados.get('id')
                        cursor.execute("""
                            INSERT INTO ai_historico (usuario, tipo, input, output, data_hora, processo_id)
                            VALUES (?, ?, ?, ?, ?, ?)
                        """, (st.session_state.user, 'sugestao', str(processo_dados), resposta, datetime.now().isoformat(), processo_id))
                        conn.commit()
                except Exception as e:
                    st.error(f"Erro ao salvar sugestões: {e}")
    
    except Exception as e:
        st.error(f"Erro ao carregar processos: {e}")


def render_historico():
    """Renderiza histórico de interações com IA"""
    st.subheader("📚 Histórico de Interações")
    
    try:
        # Buscar histórico do banco
        with db.get_connection() as conn:
            historico = pd.read_sql_query("""
                SELECT * FROM ai_historico 
                WHERE usuario = ?
                ORDER BY data_hora DESC
                LIMIT 50
            """, conn, params=(st.session_state.user,))
        
        if historico.empty:
            st.info("Nenhuma interação registrada ainda")
            return
        
        # Filtros
        col1, col2 = st.columns(2)
        with col1:
            filtro_tipo = st.multiselect(
                "Filtrar por tipo:",
                options=['chat', 'analise', 'sugestao'],
                default=['chat', 'analise', 'sugestao']
            )
        
        # Aplicar filtro
        historico_filtrado = historico[historico['tipo'].isin(filtro_tipo)]
        
        # Exibir histórico
        for idx, row in historico_filtrado.iterrows():
            with st.expander(f"🤖 {row['tipo'].upper()} - {row['data_hora'][:16]}"):
                st.markdown(f"**Entrada:**\n{row['input'][:200]}...")
                st.markdown(f"**Resposta:**\n{row['output'][:300]}...")
                
                if row['processo_id']:
                    st.caption(f"Vinculado ao processo ID: {row['processo_id']}")
    
    except Exception as e:
        st.error(f"Erro ao carregar histórico: {e}")

# --- FUNÇÕES DE CONTEXTO ---

def get_contexto_financeiro():
    """Coleta dados financeiros para análise da IA"""
    try:
        df = db.sql_get("financeiro")
        if df.empty: return "Sem dados financeiros."
        
        # Resumo
        total_entrada = df[df['tipo']=='Entrada']['valor'].sum()
        total_saida = df[df['tipo']=='Saída']['valor'].sum()
        saldo = total_entrada - total_saida
        
        # Inadimplência
        inadimplentes = df[(df['tipo']=='Entrada') & (df['status_pagamento']=='Pendente') & (df['vencimento'] < datetime.now().strftime('%Y-%m-%d'))]
        total_inad = inadimplentes['valor'].sum()
        
        return {
            "resumo_geral": {
                "total_entradas": total_entrada,
                "total_saidas": total_saida,
                "saldo": saldo,
                "inadimplencia_total": total_inad
            },
            "top_inadimplentes": inadimplentes[['descricao', 'valor', 'vencimento']].head(5).to_dict('records')
        }
    except Exception as e:
        return f"Erro ao buscar financeiro: {e}"

def get_contexto_processos():
    """Coleta processos parados há muito tempo"""
    try:
        df = db.sql_get("processos")
        if df.empty: return "Sem processos."
        
        # Filtrar ativos
        ativos = df[df['status'] == 'Ativo']
        # Simulação de "parados": sem andamento recente (idealmente cruzaria com tabela andamentos, mas vamos simplificar)
        # Vamos pegar os 5 mais antigos por data de distribuição
        antigos = ativos.sort_values('data_distribuicao').head(5)
        
        return {
            "total_ativos": len(ativos),
            "processos_antigos_atencao": antigos[['numero', 'cliente_nome', 'acao', 'data_distribuicao']].to_dict('records')
        }
    except Exception as e:
        return f"Erro ao buscar processos: {e}"

def get_contexto_propostas():
    """Coleta dados do funil de vendas"""
    try:
        df = db.sql_get("clientes")
        if df.empty: return "Sem clientes."
        
        em_negociacao = df[df['status_cliente'] == 'EM NEGOCIAÇÃO']
        total_potencial = em_negociacao['proposta_valor'].sum()
        
        return {
            "clientes_em_negociacao": len(em_negociacao),
            "valor_total_pipeline": total_potencial,
            "lista_propostas": em_negociacao[['nome', 'proposta_valor', 'status_proposta']].to_dict('records')
        }
    except Exception as e:
        return f"Erro ao buscar propostas: {e}"

def render_quick_actions():
    """Renderiza botões de ação rápida"""
    st.markdown("##### ⚡ Ações Rápidas")
    c1, c2, c3 = st.columns(3)
    
    prompt_auto = None
    contexto_auto = None
    
    if c1.button("💰 Analisar Financeiro", use_container_width=True):
        contexto_auto = get_contexto_financeiro()
        prompt_auto = "Analise a saúde financeira do escritório com base nestes dados. Identifique pontos de atenção, inadimplência e sugira melhorias."
        
    if c2.button("⚖️ Processos Parados", use_container_width=True):
        contexto_auto = get_contexto_processos()
        prompt_auto = "Analise estes processos que estão antigos ou parados. Sugira despachos ou medidas para dar andamento."
        
    if c3.button("🤝 Analisar Propostas", use_container_width=True):
        contexto_auto = get_contexto_propostas()
        prompt_auto = "Analise o funil de vendas e as propostas em aberto. Sugira estratégias para fechar esses contratos."
        
    # Se clicou em algum botão, processa automaticamente
    if prompt_auto:
        # Adicionar ao histórico visual
        st.session_state.chat_history.append({"role": "user", "content": f"🔄 [Ação Rápida] {prompt_auto}"})
        
        with st.spinner("🤖 Analisando dados do sistema..."):
            resposta = ai.chat_assistente(prompt_auto, contexto=contexto_auto)
            st.session_state.chat_history.append({"role": "assistant", "content": resposta})
            st.rerun()
