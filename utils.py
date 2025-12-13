import requests
import database as db
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import pandas as pd
import PyPDF2
from datetime import datetime, timedelta
import re

# --- HELPERS DE FORMATAÇÃO E VALIDAÇÃO ---

def limpar_numeros(valor):
    """Remove tudo que não for dígito."""
    if not valor: return ""
    return re.sub(r'\D', '', str(valor))

def safe_float(val):
    """Converte para float de forma segura."""
    try:
        if isinstance(val, (int, float)): return float(val)
        if not val: return 0.0
        val = str(val).replace('R$', '').replace('.', '').replace(',', '.').strip()
        return float(val)
    except: return 0.0

def safe_int(val):
    """Converte para int de forma segura."""
    try: return int(float(val)) if val else 1
    except: return 1

def formatar_moeda(valor):
    """Formata float para moeda BRL."""
    try:
        return f"R$ {valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except:
        return "R$ 0,00"

def formatar_documento(doc, tipo=None):
    """Formata CPF ou CNPJ."""
    if not doc:
        return ""
    d = limpar_numeros(doc)
    if len(d) == 11: # CPF
        return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}"
    elif len(d) == 14: # CNPJ
        return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}"
    return doc

def formatar_cpf(cpf):
    """Formata CPF especificamente."""
    if not cpf:
        return ""
    d = limpar_numeros(cpf)
    if len(d) == 11:
        return f"{d[:3]}.{d[3:6]}.{d[6:9]}-{d[9:]}"
    return cpf

def formatar_celular(telefone):
    """Formata número de telefone brasileiro."""
    if not telefone:
        return ""
    d = limpar_numeros(str(telefone))
    if len(d) == 11:  # Celular com DDD
        return f"({d[:2]}) {d[2:7]}-{d[7:]}"
    elif len(d) == 10:  # Fixo com DDD
        return f"({d[:2]}) {d[2:6]}-{d[6:]}"
    return str(telefone)

def formatar_cep(cep):
    """Formata CEP: 99999-999."""
    if not cep: return ""
    d = limpar_numeros(str(cep))
    if len(d) == 8:
        return f"{d[:5]}-{d[5:]}"
    return d # Retorna limpo ou original se tamanho errado

def formatar_rg(rg):
    """Formata RG (Aceita 8 ou 9 dígitos)."""
    if not rg: return ""
    d = limpar_numeros(str(rg))
    
    # Padrão SP (9 dígitos): 12.345.678-9
    if len(d) == 9:
        return f"{d[:2]}.{d[2:5]}.{d[5:8]}-{d[8]}"
    
    # Padrão RJ/Outros (8 dígitos): 12.345.678-9 ou 1.234.567-8
    # Vamos adotar XX.XXX.XXX para 8 dígitos simples ou X.XXX.XXX-X
    # Comumente 8 dígitos é: 12.345.678
    elif len(d) == 8:
        return f"{d[:2]}.{d[2:5]}.{d[5:8]}"
        
    return rg

def formatar_data(data_iso):
    """Formata data ISO (YYYY-MM-DD) para BR (DD/MM/YYYY)."""
    try:
        if not data_iso:
            return ""
        if 'T' in data_iso:
            data = datetime.fromisoformat(data_iso.replace("Z", "+00:00"))
        else:
            data = datetime.strptime(data_iso, "%Y-%m-%d")
        return data.strftime("%d/%m/%Y")
    except:
        return data_iso

def formatar_data_extenso():
    """Retorna a data atual por extenso em português."""
    meses = {
        1: 'janeiro', 2: 'fevereiro', 3: 'março', 4: 'abril',
        5: 'maio', 6: 'junho', 7: 'julho', 8: 'agosto',
        9: 'setembro', 10: 'outubro', 11: 'novembro', 12: 'dezembro'
    }
    hoje = datetime.now()
    return f"{hoje.day} de {meses[hoje.month]} de {hoje.year}"

def validar_cpf_matematico(cpf):
    """Valida CPF matematicamente."""
    cpf = limpar_numeros(cpf)
    if len(cpf) != 11 or cpf == cpf[0] * 11: return False
    soma = sum(int(cpf[i]) * (10 - i) for i in range(9))
    resto = (soma * 10) % 11
    d1 = 0 if resto == 10 else resto
    if d1 != int(cpf[9]): return False
    soma = sum(int(cpf[i]) * (11 - i) for i in range(10))
    resto = (soma * 10) % 11
    d2 = 0 if resto == 10 else resto
    return d2 == int(cpf[10])

def validar_cnpj(cnpj):
    """Valida CNPJ matematicamente."""
    cnpj = limpar_numeros(cnpj)
    if len(cnpj) != 14: return False
    
    # Validação do primeiro dígito verificador
    pesos1 = [5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    soma1 = sum(int(cnpj[i]) * pesos1[i] for i in range(12))
    resto1 = soma1 % 11
    d1 = 0 if resto1 < 2 else 11 - resto1
    if d1 != int(cnpj[12]): return False
    
    # Validação do segundo dígito verificador
    pesos2 = [6, 5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
    soma2 = sum(int(cnpj[i]) * pesos2[i] for i in range(13))
    resto2 = soma2 % 11
    d2 = 0 if resto2 < 2 else 11 - resto2
    return d2 == int(cnpj[13])

def mask_sensitive_data(text):
    """Mascara CPF, CNPJ e outros dados sensíveis em texto para conformidade LGPD."""
    if not text:
        return text
    
    text = str(text)
    
    # Mascarar CPF (123.456.789-01 -> ***.***.***-01)
    text = re.sub(r'\d{3}\.\d{3}\.\d{3}-(\d{2})', r'***.***.***-\1', text)
    
    # Mascarar CPF sem formatação (12345678901 -> ***********01)
    text = re.sub(r'(?<!\d)\d{9}(\d{2})(?!\d)', r'*********\1', text)
    
    # Mascarar CNPJ (12.345.678/0001-90 -> **.***.***/****-90)
    text = re.sub(r'\d{2}\.\d{3}\.\d{3}/\d{4}-(\d{2})', r'**.***.***/****-\1', text)
    
    # Mascarar CNPJ sem formatação
    text = re.sub(r'(?<!\d)\d{12}(\d{2})(?!\d)', r'************\1', text)
    
    # Mascarar nomes (deixa apenas primeira e última letra)
    # Exemplo: "João Silva Santos" -> "J*** S**** S*****"
    def mask_name(match):
        words = match.group(0).split()
        masked = []
        for word in words:
            if len(word) > 2:
                masked.append(word[0] + '*' * (len(word) - 1))
            else:
                masked.append(word)
        return ' '.join(masked)
    
    # Mascarar sequências de palavras capitalizadas (possíveis nomes)
    text = re.sub(r'(?:[A-ZÇÁÉÍÓÚÂÊÔÃÕ][a-zçáéíóúâêôãõ]+\s+){1,3}[A-ZÇÁÉÍÓÚÂÊÔÃÕ][a-zçáéíóúâêôãõ]+', mask_name, text)
    
    return text

def validar_email(email):
    """Valida formato de email."""
    if not email:
        return False
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return bool(re.match(pattern, str(email).strip()))

def validar_telefone(telefone):
    """Valida número de telefone brasileiro."""
    if not telefone:
        return False
    d = limpar_numeros(telefone)
    
    # DDDs válidos no Brasil (11-99, exceto alguns)
    ddds_validos = [
        '11', '12', '13', '14', '15', '16', '17', '18', '19',  # SP
        '21', '22', '24',  # RJ
        '27', '28',  # ES
        '31', '32', '33', '34', '35', '37', '38',  # MG
        '41', '42', '43', '44', '45', '46',  # PR
        '47', '48', '49',  # SC
        '51', '53', '54', '55',  # RS
        '61',  # DF
        '62', '64',  # GO
        '63',  # TO
        '65', '66',  # MT
        '67',  # MS
        '68',  # AC
        '69',  # RO
        '71', '73', '74', '75', '77',  # BA
        '79',  # SE
        '81', '87',  # PE
        '82',  # AL
        '83',  # PB
        '84',  # RN
        '85', '88',  # CE
        '86', '89',  # PI
        '91', '93', '94',  # PA
        '92', '97',  # AM
        '95',  # RR
        '96',  # AP
        '98', '99'  # MA
    ]
    
    # Deve ter 10 ou 11 dígitos
    if len(d) not in [10, 11]:
        return False
    
    # Verifica DDD
    ddd = d[:2]
    if ddd not in ddds_validos:
        return False
    
    # Se tem 11 dígitos, deve ser celular (9 XXXX-XXXX)
    if len(d) == 11:
        if d[2] != '9':
            return False
    
    # Se tem 10 dígitos, deve ser fixo (não pode começar com 9)
    if len(d) == 10:
        if d[2] == '9':
            return False
    
    return True

def buscar_cep(cep):
    """Busca endereço pelo CEP usando API ViaCEP."""
    try:
        cep_limpo = limpar_numeros(cep)
        if len(cep_limpo) != 8:
            return None
        
        url = f"https://viacep.com.br/ws/{cep_limpo}/json/"
        response = requests.get(url, timeout=5)
        
        if response.status_code == 200:
            data = response.json()
            if "erro" not in data:
                return data
        return None
    except Exception as e:
        return None

# --- CÁLCULOS ---
def calc_venc(data_pub, dias, regra="Dias Úteis"):
    """Calcula data de vencimento (fatal) baseada na regra de contagem."""
    try:
        # Garantir tipo data
        if isinstance(data_pub, str):
            data_pub = datetime.strptime(data_pub, '%Y-%m-%d').date()
        elif isinstance(data_pub, datetime):
            data_pub = data_pub.date()
            
        if regra == "Corridos":
            return data_pub + timedelta(days=dias)
        else:
            # Dias Úteis (Simples - Pula Sábado e Domingo)
            current = data_pub
            add = 0
            while add < dias:
                current += timedelta(days=1)
                if current.weekday() < 5: # 0=Seg, 4=Sex, 5=Sab, 6=Dom
                    add += 1
            return current
    except:
        return datetime.now().date()

def calcular_farol(d):
    try:
        delta = (datetime.strptime(d, '%Y-%m-%d').date() - datetime.now().date()).days
        if delta < 0: return "⚫ Vencido"
        return "🔴 Urgente" if delta <= 3 else "🟡 Atenção" if delta <= 7 else "🟢 No Prazo"
    except: return "⚪"

def gerar_documento(tipo, dados, opcoes={}):
    """Gera documento Word baseado no modelo."""
    if isinstance(dados, pd.Series):
        dados = dados.to_dict()
        
    doc = Document()
    
    # Configurar Estilo Padrão (Arial 12)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Configurar Margens
    section = doc.sections[0]
    section.top_margin = Pt(30)
    section.bottom_margin = Pt(20)
    section.left_margin = Pt(40)
    section.right_margin = Pt(20)

    # --- CABEÇALHO ---
    header = section.header
    p_header = header.paragraphs[0]
    p_header.text = ""
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tentar carregar logo
    try:
        run_img = p_header.add_run()
        run_img.add_picture('LOGO.jpg', width=Pt(150))
    except:
        pass
    
    p_header.add_run("\n")
    
    # --- PROPOSTA ---
    if tipo == "Proposta":
        t = doc.add_heading('PROPOSTA DE HONORÁRIOS ADVOCATÍCIOS', level=1)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo com Objeto Resumido
        obj_resumo = dados.get('proposta_objeto', 'Serviços Jurídicos')
        if not obj_resumo: obj_resumo = 'Serviços Jurídicos'
        obj_resumo = obj_resumo.split('\\n')[0][:50]
        
        p_sub = doc.add_paragraph(f"({obj_resumo}...)")
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y')}   Validade: 10 dias corridos")
        
        doc.add_paragraph(f"CONTRATANTE: {dados.get('nome', '').upper()}")
        
        cpf_fmt = formatar_documento(dados.get('cpf_cnpj', ''))
        tel_fmt = formatar_celular(dados.get('telefone', ''))
        doc.add_paragraph(f"CPF/CNPJ: {cpf_fmt}")
        doc.add_paragraph(f"TELEFONE: {tel_fmt}")
        
        oab = db.get_config('oab', 'OAB/RJ nº 215691')
        nome_adv = db.get_config('nome_escritorio', 'Dra. Sheila Lopes')
        doc.add_paragraph(f"CONTRATADA: {nome_adv} – {oab}")
        
        # 1. Objeto
        doc.add_heading('1. OBJETO DOS SERVIÇOS', level=2)
        obj_txt = dados.get('proposta_objeto')
        if not obj_txt: obj_txt = 'Prestação de serviços jurídicos.'
        doc.add_paragraph(obj_txt)
        
        # 2. Serviços Incluídos (Texto Padrão Profissional)
        doc.add_heading('2. SERVIÇOS INCLUÍDOS', level=2)
        p_inc = doc.add_paragraph()
        p_inc.add_run("Os serviços abrangidos por esta proposta incluem:\n").bold = True
        itens_inclusos = [
            "Reuniões e consultoria jurídica referente ao caso durante a tramitação do feito.",
            "Análise detalhada da documentação fornecida pelo Contratante.",
            "Elaboração e distribuição da Petição Inicial, com eventual pedido de Tutela de Urgência (liminar) para fixação provisória da guarda e/ou regime de convivência.",
            "Acompanhamento de todos os atos e publicações processuais em Primeira Instância.",
            "Elaboração de petições incidentais necessárias (manifestações, réplicas, etc.).",
            "Participação em audiências (conciliação, mediação e instrução).",
            "Acompanhamento de eventuais estudos psicossociais determinados pelo Juízo.",
            "Elaboração de alegações finais.",
            "Acompanhamento até a prolação da Sentença pelo Juiz de primeiro grau."
        ]
        for item in itens_inclusos:
            p_inc.add_run(f"• {item}\n")
            
        # 3. Serviços Não Incluídos
        doc.add_heading('3. SERVIÇOS NÃO INCLUÍDOS', level=2)
        p_nao_inc = doc.add_paragraph()
        p_nao_inc.add_run("Não estão contemplados nesta proposta de honorários:\n").bold = True
        itens_nao_inclusos = [
            "Acompanhamento e interposição de eventuais Recursos para instâncias superiores (Tribunal de Justiça, STJ, STF).",
            "Ações incidentais autônomas (Ex: Ação de Prestação de Contas de Alimentos, Cumprimento de Sentença, Alienação Parental em autos apartados, etc.).",
            "Custas processuais, taxas judiciárias, despesas com perícias (psicossociais, se não cobertas pela gratuidade), emolumentos de cartório, honorários de sucumbência (pagos à parte contrária em caso de derrota) e outras despesas processuais.",
            "Despesas de locomoção para atos fora da Comarca de Maricá, caso necessário."
        ]
        for item in itens_nao_inclusos:
            p_nao_inc.add_run(f"• {item}\n")
        doc.add_paragraph("Obs.: A contratação para eventuais serviços não incluídos, como Recursos, dependerá de nova proposta e contrato específico.")

        # 4. Honorários
        doc.add_heading('4. HONORÁRIOS ADVOCATÍCIOS', level=2)
        val_total = safe_float(dados.get('proposta_valor'))
        doc.add_paragraph(f"Pelos serviços jurídicos descritos, os honorários ficam ajustados no valor total de {formatar_moeda(val_total)}.")
        
        # 5. Condições de Pagamento
        doc.add_heading('5. CONDIÇÕES DE PAGAMENTO', level=2)
        val_ent = safe_float(dados.get('proposta_entrada'))
        val_saldo = val_total - val_ent
        n_parc = safe_int(dados.get('proposta_parcelas'))
        forma = dados.get('proposta_pagamento', 'A Combinar')
        
        v_parc = val_saldo / n_parc if n_parc > 0 else 0
        
        p_pag = doc.add_paragraph(f"O valor total será pago da seguinte forma ({forma}):\n")
        if val_ent > 0:
            p_pag.add_run(f"5.1. ENTRADA: {formatar_moeda(val_ent)}, no ato da assinatura.\n")
        if val_saldo > 0:
            txt_saldo = f"5.2. SALDO REMANESCENTE: {formatar_moeda(val_saldo)}, dividido em {n_parc} parcelas de {formatar_moeda(v_parc)}"
            
            # Adicionar data da primeira parcela se existir
            data_pag = dados.get('proposta_data_pagamento')
            if data_pag:
                try:
                    d_fmt = datetime.strptime(data_pag, '%Y-%m-%d').strftime('%d/%m/%Y')
                    txt_saldo += f", vencendo a primeira em {d_fmt}"
                except: pass
            
            txt_saldo += "."
            p_pag.add_run(txt_saldo)
            
        doc.add_paragraph("Obs.: O não pagamento na data aprazada implicará em multa de 2% e juros de 1% ao mês.")
        
        # 6. Sucumbência
        doc.add_heading('6. HONORÁRIOS DE SUCUMBÊNCIA', level=2)
        doc.add_paragraph("Eventuais honorários de sucumbência (valores pagos pela parte contrária em caso de êxito na ação, fixados pelo Juiz) pertencerão exclusivamente à Contratada (Advogada), conforme o Art. 23 da Lei nº 8.906/94 (Estatuto da Advocacia e da OAB), não se confundindo com os honorários contratuais aqui ajustados.")
        
        # 7. Aceite
        doc.add_heading('7. ACEITE', level=2)
        doc.add_paragraph("O aceite desta proposta se dará mediante a assinatura do respectivo Contrato de Prestação de Serviços Advocatícios, que detalhará todas as obrigações das partes, e o efetivo pagamento da Entrada (item 5.1).")
        doc.add_paragraph("Coloco-me à disposição para quaisquer esclarecimentos que se façam necessários.")
        
        doc.add_paragraph("\n\n")
        
        # Assinaturas
        tab = doc.add_table(rows=1, cols=2)
        tab.autofit = True
        c1 = tab.cell(0, 0)
        c2 = tab.cell(0, 1)
        
        p1 = c1.paragraphs[0]
        nome_adv = db.get_config('nome_escritorio', 'Dra. Sheila Lopes')
        p1.add_run(f"___________________________\n{nome_adv}\nAdvogada").bold = True
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p2 = c2.paragraphs[0]
        p2.add_run(f"___________________________\n{dados.get('nome', 'Cliente')}\nCiente e De Acordo").bold = True
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- PROCURAÇÃO ---
    elif tipo == "Procuracao":
        # Título
        p_titulo = doc.add_paragraph('PROCURAÇÃO AD JUDICIA ET EXTRA')
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.runs[0]
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)
        run_titulo.font.color.rgb = RGBColor(0, 0, 0) # Preto
        
        doc.add_paragraph("\n")
        
        # Qualificação
        end_txt = f"{dados.get('endereco', '')}, {dados.get('numero_casa', '')}, {dados.get('complemento', '')}, {dados.get('bairro', '')}, {dados.get('cidade', '')}-{dados.get('estado', '')}, CEP {dados.get('cep', '')}"
        qualif = f"OUTORGANTE: {dados.get('nome', '').upper()}, nacionalidade brasileira, {dados.get('estado_civil', '')}, {dados.get('profissao', '')}, inscrito no CPF sob nº {dados.get('cpf_cnpj', '')}, residente e domiciliado em {end_txt}."
        
        p_qualif = doc.add_paragraph(qualif)
        p_qualif.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        nome_adv = db.get_config('nome_escritorio', 'Dra. Sheila Lopes')
        oab = db.get_config('oab', 'OAB/RJ nº 215691')
        end_adv = db.get_config('endereco_escritorio', 'Rodovia Amaral Peixoto km 22, nº 5, São José do Imbassaí, Maricá/RJ')
        
        doc.add_paragraph(f"\nOUTORGADO: {nome_adv}, advogada inscrita na {oab}, com escritório profissional na {end_adv}.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph("\nPODERES: Pelo presente instrumento particular de procuração, o(a) Outorgante nomeia e constitui o(a) Outorgado(a) seu(sua) bastante procurador(a), conferindo-lhe amplos poderes para o foro em geral, com a cláusula \"ad judicia et extra\", em qualquer Juízo, Instância ou Tribunal.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Poderes Especiais (Opcional)
        if opcoes.get('poderes_especiais'):
            doc.add_paragraph("\nPODERES ESPECIAIS: Conferem-se ainda poderes específicos para receber citação, confessar, reconhecer a procedência do pedido, transigir, desistir, renunciar ao direito sobre o qual se funda a ação, receber, dar quitação, firmar compromisso e assinar declaração de hipossuficiência econômica.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        # Finalidade
        objeto = dados.get('proposta_objeto')
        if not objeto: objeto = "Ação Judicial"
        doc.add_paragraph(f"\nFINALIDADE: Especialmente para propor e acompanhar {objeto}.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph(f"\nMaricá/RJ, {formatar_data_extenso()}.").alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph("\n\n__________________________________________________")
        doc.add_paragraph(dados.get('nome', '').upper()).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- HIPOSSUFICIÊNCIA ---
    elif tipo == "Hipossuficiencia":
        # Título
        p_titulo = doc.add_paragraph('AFIRMAÇÃO DE HIPOSSUFICIÊNCIA')
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.runs[0]
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)
        run_titulo.font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph("\n") # Espaço
        
        # Dados
        nome = dados.get('nome', '').upper()
        nacionalidade = dados.get('nacionalidade', 'brasileira')
        est_civil = dados.get('estado_civil', '')
        profissao = dados.get('profissao', '')
        rg = dados.get('rg', '')
        orgao = dados.get('orgao_emissor', '')
        cpf = dados.get('cpf_cnpj', '')
        
        end_txt = f"{dados.get('endereco', '')}, nº {dados.get('numero_casa', '')}, {dados.get('complemento', '')}, {dados.get('bairro', '')}, {dados.get('cidade', '')}/{dados.get('estado', '')}"
        
        # Texto conforme imagem
        texto = f"Eu {nome}, {nacionalidade}, {profissao}, {est_civil}, portadora da cédula de identidade RG nº {rg} {orgao}, inscrita no CPF sob o nº {cpf}, residente e domiciliada na {end_txt}; afirmo, para o fim de concessão do benefício da Gratuidade de Justiça, sob as penas da lei, não possuir condições de arcar com o pagamento das custas judiciais, honorários de advogado e demais encargos, sem prejuízo de meu sustento e de minha família."
        
        p = doc.add_paragraph(texto)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(30) # Recuo primeira linha
        
        doc.add_paragraph("\n")
        
        # Data
        p_data = doc.add_paragraph(f"Maricá, {formatar_data_extenso()}.")
        p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("\n\n")
        
        # Assinatura
        p_ass = doc.add_paragraph("__________________________________________________")
        p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nome = doc.add_paragraph(nome)
        p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nome.bold = True

    # --- CONTRATO ---
    elif tipo == "Contrato":
        doc.add_heading('CONTRATO DE HONORÁRIOS', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        end_txt = f"{dados.get('endereco', '')}, {dados.get('numero_casa', '')}, {dados.get('complemento', '')}, {dados.get('bairro', '')}, {dados.get('cidade', '')}-{dados.get('estado', '')}, CEP {dados.get('cep', '')}"
        
        obj = dados.get('proposta_objeto', 'Serviços Jurídicos')
        if not obj: obj = 'Serviços Jurídicos'
        val = formatar_moeda(safe_float(dados.get('proposta_valor')))
        
        texto = f"CONTRATANTE: {dados.get('nome', '').upper()}, CPF/CNPJ {dados.get('cpf_cnpj', '')}.\nENDEREÇO: {end_txt}.\n\nOBJETO: {obj}.\nVALOR ACORDADO: {val}."
        
        p_final = doc.add_paragraph(texto)
        p_final.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph(f"\nMaricá/RJ, {formatar_data_extenso()}.\n\n")
        
        sig = doc.add_paragraph("__________________________________________________")
        sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_nm = doc.add_paragraph(dados.get('nome', 'Lopes & Ribeiro').upper())
        sig_nm.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- RODAPÉ (Igual Imagem) ---
    section = doc.sections[0]
    footer = section.footer
    p_foot = footer.paragraphs[0]
    
    end_adv = db.get_config('endereco_escritorio', 'Rodovia Amaral Peixoto km 22, nº 5, São José do Imbassaí, Maricá/RJ')
    tel_adv = db.get_config('telefone_escritorio', '(21) 97032-0748')
    email_adv = db.get_config('email_escritorio', 'sheilaadv.contato@gmail.com')
    
    p_foot.text = f"{end_adv}\n{tel_adv} | {email_adv}"
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.style.font.size = Pt(8)
    p_foot.style.font.color.rgb = None # Cinza se quisesse

    b = BytesIO(); doc.save(b); b.seek(0); return b

# Alias para compatibilidade
criar_doc = gerar_documento